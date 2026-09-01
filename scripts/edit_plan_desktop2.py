# -*- coding: utf-8 -*-
"""Desktop plan: add 5th member 庄翌 to cover; clean 6.1 team instruction."""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

SRC = r"C:\Users\34392\Desktop\千年晤旅_项目计划书_高教主赛道.docx"
doc = Document(SRC)
paras = doc.paragraphs

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)

# P14 封面团队成员：追加庄翌
t = paras[14].text
print("P14 原:", t)
if "欧阳琛斌" in t and "庄翌" not in t:
    set_text(paras[14], t + " 庄翌")
    print("P14 新:", paras[14].text)

# P106 6.1 团队构成：去掉模板提示
t = paras[106].text
print("P106 原:", t)
if "请据实填写" in t:
    set_text(paras[106], "团队共 5 人，均为在校本科生，构成互补、分工明确：成员为刘承弈、胡超悦、罗荣琪、欧阳琛斌、庄翌，具体分工见下表。")
    print("P106 新:", paras[106].text)

doc.save(SRC)
print("已保存:", SRC)
