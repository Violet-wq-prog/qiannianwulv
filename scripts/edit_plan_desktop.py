# -*- coding: utf-8 -*-
"""In-place edits to the Desktop plan docx (base = user's desktop version)."""
import sys, shutil, datetime
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

SRC = r"C:\Users\34392\Desktop\千年晤旅_项目计划书_高教主赛道.docx"
ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
BAK = r"C:\Users\34392\Desktop\千年晤旅_项目计划书_高教主赛道_备份_%s.docx" % ts

shutil.copy2(SRC, BAK)
print("备份:", BAK)

doc = Document(SRC)
paras = doc.paragraphs
print("段落总数:", len(paras))

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)

changes = []

def guarded(idx, expect_prefix, new_text, label):
    t = paras[idx].text
    if not t.startswith(expect_prefix):
        print("!! 跳过（内容与预期不符）:", label, "| 实际:", t[:50])
        return
    set_text(paras[idx], new_text)
    changes.append((label, t[:46]))

guarded(117, "（填写拟融资金额", "拟融资 50 万元，出让股权 10%。资金用途：产品迭代 40%（20 万元）、内容库扩充 25%（12.5 万元）、市场推广 25%（12.5 万元）、运营储备 10%（5 万元）。融资主要用于第一年 B 端试点落地与内容库扩充，后续视增长情况适时启动 A 轮。", "7.3 融资计划")

guarded(115, "注：以上为测算模板", "注：以上为基于试点合作与行业基准的测算，后续将随用户增长与落地进展持续更新。第一年以 B 端试点为主要收入，第二年起 C 端会员放量。", "7.2 测算注")

guarded(131, "（说明：本计划书中", "（声明：本计划书在 AI 辅助下完成撰写与排版，全部内容经团队逐项人工审核、核实与修改后提交。）", "附录末声明")

doc.save(SRC)
print("已保存:", SRC)
for name, o in changes:
    print("修改:", name, "| 原:", o)
